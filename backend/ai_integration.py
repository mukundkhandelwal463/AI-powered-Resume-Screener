import io
import json
import os
import re
import time

import docx
import google.generativeai as genai
from dotenv import load_dotenv
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fpdf import FPDF

load_dotenv()
SYSTEM_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_MODEL_CANDIDATES = (
    "gemini-2.5-flash",
    "gemini-2.0-flash",
    "gemini-flash-latest",
    "gemini-1.5-flash",
)
_RESOLVED_GEMINI_MODELS = None


def is_valid_key():
    return SYSTEM_API_KEY and SYSTEM_API_KEY.strip() != "your_gemini_api_key_here"


if is_valid_key():
    genai.configure(api_key=SYSTEM_API_KEY)


def _clip_text(text, max_chars=4000):
    txt = str(text or "")
    if len(txt) <= max_chars:
        return txt
    return txt[:max_chars]


def _generate_with_timeout(model, prompt, timeout_sec=10, retries=1):
    # Keeps API calls responsive and retries transient gateway/deadline failures.
    last_error = None
    for attempt in range(retries + 1):
        try:
            return model.generate_content(prompt, request_options={"timeout": timeout_sec}).text
        except Exception as exc:
            last_error = exc
            err = str(exc).lower()
            is_transient = any(
                token in err
                for token in ("deadline", "timeout", "timed out", "504", "503", "429", "temporar", "unavailable")
            )
            if attempt < retries and is_transient:
                time.sleep(0.8 * (attempt + 1))
                continue
            raise last_error


def _generate_with_model_fallback(prompt, timeout_sec=10):
    global _RESOLVED_GEMINI_MODELS

    if _RESOLVED_GEMINI_MODELS is None:
        discovered = []
        try:
            for model in genai.list_models():
                methods = getattr(model, "supported_generation_methods", []) or []
                if "generateContent" not in methods:
                    continue
                name = getattr(model, "name", "")
                short_name = name.split("/", 1)[-1] if "/" in name else name
                if short_name:
                    discovered.append(short_name)
        except Exception:
            discovered = []

        if discovered:
            ordered = [name for name in GEMINI_MODEL_CANDIDATES if name in discovered]
            _RESOLVED_GEMINI_MODELS = ordered or discovered[:4]
        else:
            _RESOLVED_GEMINI_MODELS = list(GEMINI_MODEL_CANDIDATES)

    last_error = None
    for model_name in _RESOLVED_GEMINI_MODELS:
        try:
            model = genai.GenerativeModel(model_name)
            return _generate_with_timeout(model, prompt, timeout_sec=timeout_sec, retries=1)
        except Exception as exc:
            last_error = exc
            err = str(exc).upper()
            if "API_KEY_INVALID" in err:
                raise
            continue
    if last_error:
        raise last_error
    raise RuntimeError("No Gemini model candidates available.")


def get_gemini_ats_feedback(resume_text, job_description, score):
    if not is_valid_key():
        return "Please add a valid Gemini API Key to backend/.env to unlock AI feedback."
    try:
        prompt = f"""You are an expert ATS and career coach.
I have a resume and a job description. The match score is {score}%.
Job Description: {_clip_text(job_description, 2500)}
Resume: {_clip_text(resume_text, 5000)}
Give EXACTLY 5 actionable bullet points to improve the resume for this job.
Format each as numbered list (1..5), concise, no extra text."""
        return _generate_with_model_fallback(prompt, timeout_sec=14)
    except Exception as e:
        if "API_KEY_INVALID" in str(e):
            return "The Gemini API key in backend/.env is invalid."
        return f"Gemini Error: {e}"


def get_gemini_career_strategy(resume_text):
    if not is_valid_key():
        return "Please add a valid Gemini API Key to backend/.env to unlock AI career strategy."
    try:
        input_type = "Job Target/Category" if len(resume_text.split()) < 30 else "Full Resume Profile"
        prompt = f"""You are an elite career strategist.
I am providing you with my {input_type}:
"{_clip_text(resume_text, 4500)}"

Please provide a highly simplified, actionable 'Career Roadmap'.
Requirements:
1. Do not use large markdown headers (like # or ##). Avoid wordy introductions.
2. Provide a single text-based visual roadmap (using simple textual arrows like '→' or steps like 'Step 1: ...') showing how to progress from this {input_type} to a senior position.
3. Keep formatting extremely simple and minimal."""
        return _generate_with_model_fallback(prompt, timeout_sec=14)
    except Exception as e:
        if "API_KEY_INVALID" in str(e):
            return "The Gemini API key in backend/.env is invalid."
        return f"Gemini Error: {e}"


def get_gemini_resume_suggestions(data):
    if not is_valid_key():
        return "Please add a valid Gemini API Key to backend/.env to unlock resume suggestions."
    try:
        prompt = f"""I am writing my resume. Based on these draft details, give 3 actionable bullet points:
Name: {data.get('name')}
Summary: {data.get('summary')}
Skills: {data.get('skills')}
Experience: {data.get('experience')}
Education: {data.get('education')}
Keep it brief and ATS-focused."""
        return _generate_with_model_fallback(prompt, timeout_sec=14)
    except Exception as e:
        if "API_KEY_INVALID" in str(e):
            return "The Gemini API key in backend/.env is invalid."
        return f"Gemini Error: {e}"

def get_gemini_enhanced_text(text=""):
    if not is_valid_key() or not text.strip():
        return text or ""
    try:
        prompt = f"""You are an elite resume editor. Rewrite the following resume experience/project description text to make it extremely professional, impactful, and ATS-optimized.
Rules:
- Keep the core meaning and facts exactly the same.
- Use strong action verbs.
- Ensure proper grammar and punctuation.
- Output ONLY the raw corrected text. Do NOT add quotes, markdown, or chatty introductions.

Text to rewrite:
{_clip_text(text, 1000)}
"""
        result = _generate_with_model_fallback(prompt, timeout_sec=10) or ""
        return result.strip()
    except Exception:
        return text

def get_gemini_professional_summary(
    professional_title="",
    stream_or_category="",
    skills="",
    experience="",
    current_summary="",
):
    if not is_valid_key():
        return "Gemini Error: Missing valid Gemini API key."
    try:
        role_hint = (professional_title or "").strip()
        category_hint = (stream_or_category or "").strip()
        prompt = f"""You are an expert resume writer and ATS optimizer.
Your objective is to generate exactly a 2-sentence professional resume summary.

If `current_summary` is provided, you MUST analyze, correct, and heavily improve it, fixing any grammatical errors while keeping its core meaning and intent, condensing or expanding it to exactly 2 highly impactful sentences. 
If no `current_summary` is provided, generate a new one from scratch based on the skills and experience.

Rules:
- Write EXACTLY 2 sentences. No more, no less.
- Use a confident third-person style (no 'I', no 'my', no 'we').
- Start with the strongest role anchor:
  1) professional_title if provided
  2) else stream_or_category
- Build an ATS-friendly paragraph incorporating relevant skills naturally.
- Output plain text only, no bullets, no markdown, no quotes.

professional_title: {role_hint}
stream_or_category: {category_hint}
skills: {_clip_text(skills, 1200)}
experience: {_clip_text(experience, 2200)}
current_summary: {_clip_text(current_summary, 1200)}
"""
        text = _generate_with_model_fallback(prompt, timeout_sec=16) or ""
        text = re.sub(r"\s+", " ", str(text)).strip()
        return text
    except Exception as e:
        if "API_KEY_INVALID" in str(e):
            return "Gemini Error: The Gemini API key in backend/.env is invalid."
        return f"Gemini Error: {e}"


def get_gemini_full_resume_analysis(resume_text, job_description="", stream_or_category=""):
    fallback = {
        "category": (stream_or_category or "").strip() or "General",
        "ats_score": 0.0,
        "missing_keywords": [],
        "suggestions": ["Could not fetch full Gemini analysis right now. Please try again."],
        "gemini_ok": False,
    }
    if not is_valid_key():
        fallback["suggestions"] = ["Add a valid Gemini API key in backend/.env to enable AI analysis."]
        return fallback

    try:
        prompt = f"""You are an ATS and resume analyzer.
Return ONLY valid JSON (no markdown) with keys:
category (string), ats_score (number 0-100), missing_keywords (array of strings), suggestions (array of strings).
Rules:
- If stream_or_category is provided, use it as category unless resume strongly contradicts it.
- If job_description is empty, estimate ATS based on resume quality and relevance to stream_or_category.
- suggestions must be practical and ATS-focused, maximum 7 items.
stream_or_category: {stream_or_category}
job_description: {_clip_text(job_description, 2500)}
resume_text: {_clip_text(resume_text, 5000)}
"""
        raw = _generate_with_model_fallback(prompt, timeout_sec=16) or ""
        match = re.search(r"\{[\s\S]*\}", raw)
        if not match:
            return fallback

        parsed = json.loads(match.group(0))
        category = str(parsed.get("category", fallback["category"])).strip() or fallback["category"]
        try:
            ats_score = float(parsed.get("ats_score", 0.0))
        except Exception:
            ats_score = 0.0
        ats_score = round(max(0.0, min(100.0, ats_score)), 2)

        missing_keywords = parsed.get("missing_keywords", [])
        if not isinstance(missing_keywords, list):
            missing_keywords = []
        missing_keywords = [str(x).strip() for x in missing_keywords if str(x).strip()][:15]

        suggestions = parsed.get("suggestions", [])
        if not isinstance(suggestions, list):
            suggestions = []
        suggestions = [str(x).strip() for x in suggestions if str(x).strip()][:10]
        if not suggestions:
            suggestions = fallback["suggestions"]

        return {
            "category": category,
            "ats_score": ats_score,
            "missing_keywords": missing_keywords,
            "suggestions": suggestions,
            "gemini_ok": True,
        }
    except Exception:
        return fallback


class PDFResume(FPDF):
    def footer(self):
        self.set_y(-15)
        self.set_font("Arial", "I", 8)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")


def _contact_location(data):
    return data.get("location", "") or data.get("linkedin", "")


def _split_non_empty_lines(text):
    return [line.strip() for line in str(text or "").splitlines() if line.strip()]


def _split_entry_blocks(text):
    raw = str(text or "").strip()
    if not raw:
        return []
    chunks = re.split(r"\n\s*\n", raw)
    blocks = []
    for chunk in chunks:
        lines = _split_non_empty_lines(chunk)
        if lines:
            blocks.append(lines)
    return blocks


def _split_skills(text):
    parts = re.split(r"[,\n]", str(text or ""))
    return [item.strip() for item in parts if item.strip()]


def _template_key(template):
    return str(template or "").strip().lower()


def _entry_records(text, default_label):
    records = []
    for idx, lines in enumerate(_split_entry_blocks(text), start=1):
        if len(lines) == 1:
            records.append(
                {
                    "meta": f"{default_label} {idx}",
                    "title": lines[0],
                    "details": [],
                }
            )
        elif len(lines) == 2:
            records.append(
                {
                    "meta": lines[0],
                    "title": lines[1],
                    "details": [],
                }
            )
        else:
            records.append(
                {
                    "meta": lines[0],
                    "title": lines[1],
                    "details": lines[2:],
                }
            )
    return records


def generate_pdf(data, template):
    pdf = PDFResume()
    pdf.add_page()
    location = _contact_location(data)
    template_key = _template_key(template)

    name = (data.get("name", "") or "YOUR NAME").strip()
    role = (data.get("headline", "") or "Professional Role").strip()
    summary = (data.get("summary", "") or "").strip()
    website = (data.get("website", "") or "").strip()

    def add_rule(y=None):
        line_y = y if y is not None else pdf.get_y()
        pdf.set_draw_color(140, 140, 140)
        pdf.set_line_width(0.5)
        pdf.line(14, line_y, 196, line_y)

    def render_records(x, y, width, records, empty_hint):
        curr_y = y
        if not records:
            pdf.set_xy(x, curr_y)
            pdf.set_font("Arial", "I", 10)
            pdf.set_text_color(105, 105, 105)
            pdf.multi_cell(width, 5.5, empty_hint)
            return pdf.get_y() + 2

        for record in records:
            pdf.set_xy(x, curr_y)
            pdf.set_font("Arial", "", 10)
            pdf.set_text_color(95, 95, 95)
            pdf.multi_cell(width, 5.2, record["meta"])
            curr_y = pdf.get_y()

            pdf.set_xy(x, curr_y)
            pdf.set_font("Arial", "B", 11)
            pdf.set_text_color(40, 40, 40)
            pdf.multi_cell(width, 5.6, record["title"])
            curr_y = pdf.get_y()

            if record["details"]:
                for line in record["details"]:
                    pdf.set_xy(x, curr_y)
                    pdf.set_font("Arial", "", 10)
                    pdf.set_text_color(45, 45, 45)
                    pdf.multi_cell(width, 5.3, f"- {line}")
                    curr_y = pdf.get_y()

            curr_y += 2
        return curr_y

    if template_key in {"classical.pdf", "resume for experienced.pdf"}:
        pdf.set_fill_color(235, 235, 235)
        pdf.rect(8, 8, 194, 281, "F")

        pdf.set_xy(10, 16)
        pdf.set_text_color(45, 45, 45)
        pdf.set_font("Arial", "B", 28)
        pdf.cell(0, 11, name.upper(), align="C", ln=True)
        pdf.set_font("Arial", "", 14)
        pdf.cell(0, 8, role, align="C", ln=True)

        pdf.set_font("Arial", "", 10.5)
        pdf.set_text_color(68, 68, 68)
        contacts = [data.get("phone", ""), data.get("email", ""), location]
        contacts = [c for c in contacts if c]
        pdf.cell(0, 7, "  |  ".join(contacts) if contacts else "+91 XXXXX XXXXX | hello@email.com | City, Country", align="C", ln=True)

        add_rule(pdf.get_y() + 1)
        pdf.set_y(pdf.get_y() + 6)

        pdf.set_font("Arial", "B", 13)
        pdf.set_text_color(40, 40, 40)
        pdf.cell(0, 7, "ABOUT ME", ln=True)
        pdf.set_font("Arial", "", 10.5)
        pdf.set_text_color(50, 50, 50)
        pdf.multi_cell(180, 6, summary or "Write a short profile about your background and strengths.")

        add_rule(pdf.get_y() + 2)
        pdf.set_y(pdf.get_y() + 7)

        pdf.set_font("Arial", "B", 13)
        pdf.set_text_color(40, 40, 40)
        pdf.cell(0, 7, "EDUCATION", ln=True)
        edu_y = render_records(14, pdf.get_y() + 1, 180, _entry_records(data.get("education", ""), "Education"), "Add education entries.")
        pdf.set_y(edu_y)

        add_rule(pdf.get_y() + 1)
        pdf.set_y(pdf.get_y() + 6)

        pdf.set_font("Arial", "B", 13)
        pdf.set_text_color(40, 40, 40)
        pdf.cell(0, 7, "WORK EXPERIENCE", ln=True)
        exp_y = render_records(14, pdf.get_y() + 1, 180, _entry_records(data.get("experience", ""), "Experience"), "Add work experience entries.")
        pdf.set_y(exp_y)

        add_rule(pdf.get_y() + 1)
        pdf.set_y(pdf.get_y() + 6)

        pdf.set_font("Arial", "B", 13)
        pdf.set_text_color(40, 40, 40)
        pdf.cell(0, 7, "SKILLS", ln=True)
        skills = _split_skills(data.get("skills", ""))
        if skills:
            pdf.set_font("Arial", "", 10.5)
            pdf.set_text_color(45, 45, 45)
            for skill in skills:
                pdf.multi_cell(180, 5.5, f"- {skill}")
        else:
            pdf.set_font("Arial", "I", 10)
            pdf.set_text_color(100, 100, 100)
            pdf.multi_cell(180, 5.5, "Add your strongest skills, separated by commas.")

    elif template_key == "freasher.pdf":
        left_x = 14
        left_w = 116
        right_x = 138
        right_w = 58

        pdf.set_fill_color(238, 238, 238)
        pdf.rect(8, 8, 194, 281, "F")

        pdf.set_xy(left_x, 16)
        pdf.set_font("Arial", "B", 24)
        pdf.set_text_color(44, 44, 44)
        pdf.multi_cell(left_w + right_w + 8, 10, name.upper())

        pdf.set_xy(left_x, 28)
        pdf.set_font("Arial", "", 13)
        pdf.set_text_color(72, 72, 72)
        pdf.multi_cell(left_w + right_w + 8, 7, role)

        add_rule(36)
        y = 40

        pdf.set_xy(left_x, y)
        pdf.set_font("Arial", "B", 13)
        pdf.set_text_color(40, 40, 40)
        pdf.multi_cell(left_w + right_w + 8, 6.5, "ABOUT ME")
        pdf.set_xy(left_x, pdf.get_y())
        pdf.set_font("Arial", "", 10.2)
        pdf.set_text_color(55, 55, 55)
        pdf.multi_cell(left_w + right_w + 8, 5.5, summary or "Write your introduction summary.")

        add_rule(pdf.get_y() + 2)
        top_cols_y = pdf.get_y() + 5

        y_left = top_cols_y
        pdf.set_xy(left_x, y_left)
        pdf.set_font("Arial", "B", 13)
        pdf.set_text_color(40, 40, 40)
        pdf.multi_cell(left_w, 6.5, "EDUCATION")
        y_left = render_records(left_x, pdf.get_y(), left_w, _entry_records(data.get("education", ""), "Education"), "Add education entries.")
        add_rule(y_left)
        y_left += 4

        pdf.set_xy(left_x, y_left)
        pdf.set_font("Arial", "B", 13)
        pdf.set_text_color(40, 40, 40)
        pdf.multi_cell(left_w, 6.5, "WORK EXPERIENCE")
        y_left = render_records(left_x, pdf.get_y(), left_w, _entry_records(data.get("experience", ""), "Experience"), "Add work experience entries.")

        y_right = top_cols_y
        pdf.set_xy(right_x, y_right)
        pdf.set_font("Arial", "B", 13)
        pdf.set_text_color(40, 40, 40)
        pdf.multi_cell(right_w, 6.5, "CONTACT")
        contact_lines = [data.get("phone", ""), data.get("email", ""), location, website]
        contact_lines = [line for line in contact_lines if line]
        if not contact_lines:
            contact_lines = ["Add contact details."]
        pdf.set_font("Arial", "", 10)
        pdf.set_text_color(55, 55, 55)
        for line in contact_lines:
            pdf.set_x(right_x)
            pdf.multi_cell(right_w, 5.3, line)
        y_right = pdf.get_y() + 2
        add_rule(y_right)
        y_right += 4

        pdf.set_xy(right_x, y_right)
        pdf.set_font("Arial", "B", 13)
        pdf.set_text_color(40, 40, 40)
        pdf.multi_cell(right_w, 6.5, "SKILLS")
        prof_skills = _split_skills(data.get("skills", ""))
        personal_skills = _split_skills(data.get("side_skills", ""))

        pdf.set_font("Arial", "B", 10)
        pdf.set_x(right_x)
        pdf.multi_cell(right_w, 5, "Professional")
        pdf.set_font("Arial", "", 10)
        for s in (prof_skills or ["Add professional skills."]):
            pdf.set_x(right_x)
            pdf.multi_cell(right_w, 5.1, f"- {s}")

        pdf.set_font("Arial", "B", 10)
        pdf.set_x(right_x)
        pdf.multi_cell(right_w, 5, "Personal")
        pdf.set_font("Arial", "", 10)
        for s in (personal_skills or ["Add personal skills."]):
            pdf.set_x(right_x)
            pdf.multi_cell(right_w, 5.1, f"- {s}")
        y_right = pdf.get_y() + 2
        add_rule(y_right)
        y_right += 4

        pdf.set_xy(right_x, y_right)
        pdf.set_font("Arial", "B", 13)
        pdf.multi_cell(right_w, 6.5, "LANGUAGES")
        languages = _split_skills(data.get("languages", ""))
        pdf.set_font("Arial", "", 10)
        for lang in (languages or ["Add language proficiency entries."]):
            pdf.set_x(right_x)
            pdf.multi_cell(right_w, 5.1, f"- {lang}")

        final_y = max(y_left, pdf.get_y())
        pdf.set_y(final_y)

    elif template_key == "resume for experienced2.pdf":
        pdf.set_fill_color(238, 238, 238)
        pdf.rect(8, 8, 194, 281, "F")

        pdf.set_xy(10, 14)
        pdf.set_font("Arial", "B", 24)
        pdf.set_text_color(44, 44, 44)
        pdf.cell(0, 10, name.upper(), align="C", ln=True)

        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 6, role.upper(), align="C", ln=True)

        contact_line = " | ".join([v for v in [location, data.get("phone", ""), data.get("email", "")] if v])
        pdf.set_font("Arial", "", 10)
        pdf.set_text_color(70, 70, 70)
        pdf.cell(0, 5.5, contact_line or "City, Country | +91 XXXXX XXXXX | hello@email.com", align="C", ln=True)
        pdf.cell(0, 5.5, website or "www.yourwebsite.com", align="C", ln=True)

        add_rule(pdf.get_y() + 2)
        pdf.set_y(pdf.get_y() + 8)

        pdf.set_font("Arial", "B", 13)
        pdf.set_text_color(40, 40, 40)
        pdf.cell(0, 7, "SUMMARY", ln=True)
        pdf.set_font("Arial", "", 10.2)
        pdf.set_text_color(55, 55, 55)
        pdf.multi_cell(180, 5.4, summary or "Write your executive summary.")

        add_rule(pdf.get_y() + 2)
        pdf.set_y(pdf.get_y() + 6)

        pdf.set_font("Arial", "B", 13)
        pdf.set_text_color(40, 40, 40)
        pdf.cell(0, 7, "WORK EXPERIENCE", ln=True)
        y_exp = render_records(14, pdf.get_y() + 1, 180, _entry_records(data.get("experience", ""), "Experience"), "Add work experience entries.")
        pdf.set_y(y_exp)

        add_rule(pdf.get_y() + 2)
        pdf.set_y(pdf.get_y() + 6)

        pdf.set_font("Arial", "B", 13)
        pdf.set_text_color(40, 40, 40)
        pdf.cell(0, 7, "EDUCATION", ln=True)
        y_edu = render_records(14, pdf.get_y() + 1, 180, _entry_records(data.get("education", ""), "Education"), "Add education entries.")
        pdf.set_y(y_edu)

    else:
        pdf.set_font("Arial", "B", 24)
        pdf.cell(0, 10, name, align="C", ln=True)
        pdf.set_font("Arial", "", 11)
        pdf.cell(0, 6, f"{data.get('email', '')} | {data.get('phone', '')} | {location}", align="C", ln=True)
        pdf.ln(5)
        for title, value in [
            ("Professional Summary", data.get("summary", "")),
            ("Key Skills", data.get("skills", "")),
            ("Experience", data.get("experience", "")),
            ("Education", data.get("education", "")),
        ]:
            pdf.set_font("Arial", "B", 13)
            pdf.cell(0, 7, title.upper(), ln=True)
            pdf.set_font("Arial", "", 10.5)
            pdf.multi_cell(180, 5.5, str(value or f"Add {title.lower()} details."))
            pdf.ln(1.5)

    return bytes(pdf.output(dest="S"))


def generate_docx(data, template):
    doc = docx.Document()
    location = _contact_location(data)
    template_key = _template_key(template)

    def add_heading(title, level_size=13):
        heading = doc.add_paragraph()
        run = heading.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(level_size)

    def add_rule():
        rule = doc.add_paragraph("_" * 96)
        rule.runs[0].font.size = Pt(7)

    def add_records_to_container(container, records, empty_text):
        if not records:
            hint = container.add_paragraph(empty_text)
            if hint.runs:
                hint.runs[0].italic = True
            return

        for record in records:
            meta_pg = container.add_paragraph(record["meta"])
            if meta_pg.runs:
                meta_pg.runs[0].font.size = Pt(10)
                meta_pg.runs[0].italic = True

            title_pg = container.add_paragraph(record["title"])
            if title_pg.runs:
                title_pg.runs[0].font.size = Pt(11)
                title_pg.runs[0].bold = True

            for detail in record["details"]:
                container.add_paragraph(detail, style="List Bullet")

    if template_key in {"classical.pdf", "resume for experienced.pdf"}:
        name_pg = doc.add_paragraph()
        name_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_pg.add_run((data.get("name", "") or "YOUR NAME").upper())
        name_run.bold = True
        name_run.font.size = Pt(28)

        role_pg = doc.add_paragraph(data.get("headline", "") or "Professional Role")
        role_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if role_pg.runs:
            role_pg.runs[0].font.size = Pt(14)

        contact_line = " | ".join([v for v in [data.get("phone", ""), data.get("email", ""), location] if v])
        contact_pg = doc.add_paragraph(contact_line or "+91 XXXXX XXXXX | hello@email.com | City, Country")
        contact_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_rule()

        add_heading("About Me")
        doc.add_paragraph(data.get("summary", "") or "Write a short profile about your background and strengths.")
        add_rule()

        add_heading("Education")
        add_records_to_container(doc, _entry_records(data.get("education", ""), "Education"), "Add education entries.")
        add_rule()

        add_heading("Work Experience")
        add_records_to_container(doc, _entry_records(data.get("experience", ""), "Experience"), "Add work experience entries.")
        add_rule()

        add_heading("Skills")
        for skill in (_split_skills(data.get("skills", "")) or ["Add your strongest skills, separated by commas."]):
            doc.add_paragraph(skill, style="List Bullet")

    elif template_key == "freasher.pdf":
        name_pg = doc.add_paragraph()
        name_pg.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_run = name_pg.add_run((data.get("name", "") or "YOUR NAME").upper())
        name_run.bold = True
        name_run.font.size = Pt(24)

        role_pg = doc.add_paragraph(data.get("headline", "") or "Professional Role")
        role_pg.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_rule()

        add_heading("About Me")
        doc.add_paragraph(data.get("summary", "") or "Write your introduction summary.")
        add_rule()

        table = doc.add_table(rows=1, cols=2)
        left = table.rows[0].cells[0]
        right = table.rows[0].cells[1]

        left_heading = left.paragraphs[0].add_run("EDUCATION")
        left_heading.bold = True
        add_records_to_container(left, _entry_records(data.get("education", ""), "Education"), "Add education entries.")

        left.add_paragraph()
        left_work = left.add_paragraph().add_run("WORK EXPERIENCE")
        left_work.bold = True
        add_records_to_container(left, _entry_records(data.get("experience", ""), "Experience"), "Add work experience entries.")

        right_heading = right.paragraphs[0].add_run("CONTACT")
        right_heading.bold = True
        for line in [data.get("phone", ""), data.get("email", ""), location, data.get("website", "")]:
            if line:
                right.add_paragraph(line)

        right.add_paragraph()
        right_skill_heading = right.add_paragraph().add_run("SKILLS")
        right_skill_heading.bold = True
        right_prof = right.add_paragraph().add_run("Professional")
        right_prof.bold = True
        for skill in (_split_skills(data.get("skills", "")) or ["Add professional skills."]):
            right.add_paragraph(skill, style="List Bullet")

        right_personal = right.add_paragraph().add_run("Personal")
        right_personal.bold = True
        for skill in (_split_skills(data.get("side_skills", "")) or ["Add personal skills."]):
            right.add_paragraph(skill, style="List Bullet")

        right.add_paragraph()
        right_lang_heading = right.add_paragraph().add_run("LANGUAGES")
        right_lang_heading.bold = True
        for lang in (_split_skills(data.get("languages", "")) or ["Add language proficiency entries."]):
            right.add_paragraph(lang, style="List Bullet")

    elif template_key == "resume for experienced2.pdf":
        name_pg = doc.add_paragraph()
        name_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_pg.add_run((data.get("name", "") or "YOUR NAME").upper())
        name_run.bold = True
        name_run.font.size = Pt(24)

        role_pg = doc.add_paragraph((data.get("headline", "") or "Business Consultant").upper())
        role_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if role_pg.runs:
            role_pg.runs[0].bold = True

        contact_line = " | ".join([v for v in [location, data.get("phone", ""), data.get("email", "")] if v])
        doc.add_paragraph(contact_line or "City, Country | +91 XXXXX XXXXX | hello@email.com").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(data.get("website", "") or "www.yourwebsite.com").alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_rule()

        add_heading("Summary")
        doc.add_paragraph(data.get("summary", "") or "Write your executive summary.")
        add_rule()

        add_heading("Work Experience")
        add_records_to_container(doc, _entry_records(data.get("experience", ""), "Experience"), "Add work experience entries.")
        add_rule()

        add_heading("Education")
        add_records_to_container(doc, _entry_records(data.get("education", ""), "Education"), "Add education entries.")

    else:
        name_pg = doc.add_paragraph((data.get("name", "") or "Candidate").upper())
        name_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if name_pg.runs:
            name_pg.runs[0].bold = True
            name_pg.runs[0].font.size = Pt(24)

        doc.add_paragraph(f"{data.get('email', '')} | {data.get('phone', '')} | {location}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_rule()

        for title, value in [
            ("Professional Summary", data.get("summary", "")),
            ("Key Skills", data.get("skills", "")),
            ("Experience", data.get("experience", "")),
            ("Education", data.get("education", "")),
        ]:
            add_heading(title)
            doc.add_paragraph(str(value or f"Add {title.lower()} details."))

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def generate_docx_from_builder(data):
    """Generate a properly formatted DOCX from the React Resume Builder data structure.

    Expects data with keys:
      personal_info: {full_name, email, phone, location, linkedin, website}
      professional_summary: str
      experience: [{position, company, start_date, end_date, is_current, description}]
      education: [{degree, field, institution, graduation_date, gpa}]
      project: [{name, description}]
      skills: [str]
      keywords: [str]
    """
    from docx.shared import Inches, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = docx.Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(1)
    style.paragraph_format.space_before = Pt(0)

    pi = data.get("personal_info", {}) or {}
    accent_hex = data.get("accent_color", "2563eb") or "2563eb"
    accent_hex = accent_hex.lstrip("#")
    if len(accent_hex) != 6:
        accent_hex = "2563eb"
    try:
        accent = RGBColor(int(accent_hex[0:2], 16), int(accent_hex[2:4], 16), int(accent_hex[4:6], 16))
    except Exception:
        accent = RGBColor(0x25, 0x63, 0xEB)

    def add_horizontal_rule():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), accent_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_section_heading(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = accent
        add_horizontal_rule()

    def format_date(date_str):
        if not date_str:
            return ""
        try:
            parts = date_str.split("-")
            from datetime import datetime
            dt = datetime(int(parts[0]), int(parts[1]), 1)
            return dt.strftime("%b %Y")
        except Exception:
            return date_str

    # ── NAME (centered, large, accent color) ──
    name_pg = doc.add_paragraph()
    name_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_pg.paragraph_format.space_after = Pt(1)
    name_run = name_pg.add_run(pi.get("full_name", "") or "Your Name")
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = accent

    # ── CONTACT LINE (centered) ──
    contact_parts = []
    if pi.get("email"):
        contact_parts.append(pi["email"])
    if pi.get("phone"):
        contact_parts.append(pi["phone"])
    if pi.get("location"):
        contact_parts.append(pi["location"])
    if pi.get("linkedin"):
        contact_parts.append(pi["linkedin"])
    if pi.get("website"):
        contact_parts.append(pi["website"])

    if contact_parts:
        contact_pg = doc.add_paragraph()
        contact_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_pg.paragraph_format.space_after = Pt(1)
        contact_run = contact_pg.add_run("  |  ".join(contact_parts))
        contact_run.font.size = Pt(8)
        contact_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    add_horizontal_rule()

    # ── PROFESSIONAL SUMMARY ──
    summary = data.get("professional_summary", "")
    if summary:
        add_section_heading("Professional Summary")
        p = doc.add_paragraph(summary)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # ── EXPERIENCE ──
    experience = data.get("experience", []) or []
    if experience:
        add_section_heading("Professional Experience")
        for exp in experience:
            # Row: Position | Dates
            pos_pg = doc.add_paragraph()
            pos_pg.paragraph_format.space_before = Pt(2)
            pos_pg.paragraph_format.space_after = Pt(0)
            pos_run = pos_pg.add_run(exp.get("position", "") or "Position")
            pos_run.bold = True
            pos_run.font.size = Pt(9)
            date_start = format_date(exp.get("start_date", ""))
            date_end = "Present" if exp.get("is_current") else format_date(exp.get("end_date", ""))
            date_str = f"{date_start} – {date_end}" if date_start else ""
            if date_str:
                pos_run2 = pos_pg.add_run(f"    |    {date_str}")
                pos_run2.font.size = Pt(8)
                pos_run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            # Company name
            company = exp.get("company", "")
            if company:
                comp_pg = doc.add_paragraph()
                comp_pg.paragraph_format.space_after = Pt(1)
                comp_run = comp_pg.add_run(company)
                comp_run.font.size = Pt(8)
                comp_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                comp_run.italic = True

            # Description as bullet points — max 4 bullets, max 160 chars each
            desc = exp.get("description", "")
            if desc:
                bullet_count = 0
                for line in desc.strip().split("\n"):
                    if bullet_count >= 4:
                        break
                    line = line.strip().lstrip("•-").strip()
                    if line:
                        # Truncate to ~160 chars (approx 2 lines)
                        if len(line) > 160:
                            line = line[:157].rstrip() + "..."
                        bp = doc.add_paragraph(line, style="List Bullet")
                        bp.paragraph_format.space_after = Pt(0)
                        for r in bp.runs:
                            r.font.size = Pt(8)
                        bullet_count += 1

    # ── PROJECTS ──
    projects = data.get("project", []) or []
    if projects:
        add_section_heading("Projects")
        for proj in projects[:3]:  # Max 3 projects
            p_name = doc.add_paragraph()
            p_name.paragraph_format.space_before = Pt(1)
            p_name.paragraph_format.space_after = Pt(0)
            run = p_name.add_run(proj.get("name", "") or "Project Name")
            run.bold = True
            run.font.size = Pt(9)
            # Show date next to project name
            proj_date = proj.get("date", "")
            if proj_date:
                date_run = p_name.add_run(f"  |  {proj_date}")
                date_run.font.size = Pt(8)
                date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            # Show link if available
            proj_link = proj.get("link", "")
            if proj_link:
                link_run = p_name.add_run(f"  •  {proj_link}")
                link_run.font.size = Pt(7)
                link_run.font.color.rgb = RGBColor(0x66, 0x88, 0xCC)

            p_desc = proj.get("description", "")
            if p_desc:
                # Truncate project description to ~200 chars
                if len(p_desc) > 200:
                    p_desc = p_desc[:197].rstrip() + "..."
                dp = doc.add_paragraph(p_desc)
                dp.paragraph_format.space_after = Pt(1)
                for r in dp.runs:
                    r.font.size = Pt(8)
                    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # ── EDUCATION ──
    education = data.get("education", []) or []
    if education:
        add_section_heading("Education")
        for edu in education:
            edu_pg = doc.add_paragraph()
            edu_pg.paragraph_format.space_before = Pt(2)
            edu_pg.paragraph_format.space_after = Pt(0)
            degree_text = edu.get("degree", "") or ""
            field = edu.get("field", "")
            if field:
                degree_text += f" in {field}"
            run = edu_pg.add_run(degree_text or "Degree")
            run.bold = True
            run.font.size = Pt(9)

            grad_date = format_date(edu.get("graduation_date", ""))
            if grad_date:
                run2 = edu_pg.add_run(f"    |    {grad_date}")
                run2.font.size = Pt(8)
                run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            inst = edu.get("institution", "")
            if inst:
                inst_pg = doc.add_paragraph()
                inst_pg.paragraph_format.space_after = Pt(0)
                inst_run = inst_pg.add_run(inst)
                inst_run.font.size = Pt(8)
                inst_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

            gpa = edu.get("gpa", "")
            if gpa:
                gpa_pg = doc.add_paragraph()
                gpa_pg.paragraph_format.space_after = Pt(1)
                gpa_run = gpa_pg.add_run(f"GPA: {gpa}")
                gpa_run.font.size = Pt(7)
                gpa_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── CERTIFICATIONS ──
    certifications = data.get("certifications", []) or []
    if certifications:
        add_section_heading("Certifications")
        for cert in certifications[:4]:  # Max 4
            cert_name = cert.get("name", "") or ""
            cert_date = cert.get("date", "") or ""
            cert_link = cert.get("link", "") or ""
            line_parts = [cert_name]
            if cert_date:
                line_parts.append(cert_date)
            if cert_link:
                line_parts.append(cert_link)
            cert_pg = doc.add_paragraph()
            cert_pg.paragraph_format.space_before = Pt(1)
            cert_pg.paragraph_format.space_after = Pt(0)
            cert_run = cert_pg.add_run(cert_name)
            cert_run.bold = True
            cert_run.font.size = Pt(9)
            if cert_date or cert_link:
                extra = []
                if cert_date: extra.append(cert_date)
                if cert_link: extra.append(cert_link)
                meta_run = cert_pg.add_run("  |  " + "  |  ".join(extra))
                meta_run.font.size = Pt(8)
                meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── CORE SKILLS (merged with keywords) ──
    skills = data.get("skills", []) or []
    keywords = data.get("keywords", []) or []
    all_skills = list(skills) + [k for k in keywords if k not in skills]
    if all_skills:
        add_section_heading("Core Skills")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run("  •  ".join(all_skills))
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
