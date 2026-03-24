import streamlit as st
import google.generativeai as genai
import io
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

st.set_page_config(page_title="Resume Maker (Model 3)", page_icon="📝", layout="wide")

# -------- PDF GENERATION --------
class PDFResume(FPDF):
    def header(self):
        pass
    def footer(self):
        self.set_y(-15)
        self.set_font("Arial", "I", 8)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")

def create_raw_pdf(data, template):
    pdf = PDFResume()
    pdf.add_page()
    
    if template == "Classic ATS":
        pdf.set_font("Arial", "B", 24)
        pdf.cell(0, 10, data['name'], align="C", ln=True)
        pdf.set_font("Arial", "", 11)
        pdf.cell(0, 6, f"{data['email']} | {data['phone']} | {data['linkedin']}", align="C", ln=True)
        pdf.ln(5)
        
    elif template == "Modern Minimalist":
        pdf.set_font("Arial", "B", 26)
        pdf.cell(0, 10, data['name'].upper(), align="L", ln=True)
        pdf.set_font("Arial", "I", 11)
        pdf.set_text_color(100, 100, 100)
        pdf.cell(0, 6, f"{data['email']}    |    {data['phone']}    |    {data['linkedin']}", align="L", ln=True)
        pdf.set_text_color(0, 0, 0)
        pdf.line(10, pdf.get_y()+2, 200, pdf.get_y()+2)
        pdf.ln(8)
        
    else: # Executive
        pdf.set_font("Times", "B", 22)
        pdf.cell(0, 8, data['name'], align="R", ln=True)
        pdf.set_font("Times", "", 12)
        pdf.cell(0, 6, f"{data['email']} | {data['phone']}", align="R", ln=True)
        pdf.cell(0, 6, data['linkedin'], align="R", ln=True)
        pdf.line(10, pdf.get_y()+2, 200, pdf.get_y()+2)
        pdf.ln(6)

    # Content loop helper
    def add_section(title, text):
        if not text.strip(): return
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 8, title.upper(), ln=True)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(2)
        pdf.set_font("Arial", "", 11)
        pdf.multi_cell(0, 6, text)
        pdf.ln(5)

    add_section("Professional Summary", data['summary'])
    add_section("Key Skills", data['skills'])
    add_section("Experience", data['experience'])
    add_section("Education", data['education'])
    
    return pdf.output(dest='S') # returns byte string

# -------- DOCX GENERATION --------
def create_raw_docx(data, template):
    doc = docx.Document()
    
    # Simple templates handling inside docx
    name_pg = doc.add_paragraph()
    if template == "Classic ATS":
        name_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif template == "Executive":
        name_pg.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    name_run = name_pg.add_run(data['name'].upper() if template == "Modern Minimalist" else data['name'])
    name_run.bold = True
    name_run.font.size = Pt(24)
    
    contact_pg = doc.add_paragraph()
    if template == "Classic ATS":
        contact_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif template == "Executive":
        contact_pg.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
    contact_pg.add_run(f"{data['email']} | {data['phone']} | {data['linkedin']}")
    
    doc.add_paragraph("_" * 75) # separator
    
    def add_docx_section(title, text):
        if not text.strip(): return
        h = doc.add_heading(title.upper(), level=1)
        p = doc.add_paragraph(text)
    
    add_docx_section("Professional Summary", data['summary'])
    add_docx_section("Key Skills", data['skills'])
    add_docx_section("Experience", data['experience'])
    add_docx_section("Education", data['education'])
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# -------- GEMINI API --------
def get_resume_suggestions(gemini_key, data):
    try:
        genai.configure(api_key=gemini_key)
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = f"""
        I am writing my resume. Based on these draft details, give me 3 highly actionable bullet points on how to improve my resume before I generate it:
        Name: {data['name']}
        Summary: {data['summary']}
        Skills: {data['skills']}
        Experience: {data['experience']}
        Education: {data['education']}
        
        Keep it brief and focus on ATS optimization and phrasing improvements.
        """
        resp = model.generate_content(prompt)
        return resp.text
    except Exception as e:
        return f"Could not generate suggestions. Ensure your Gemini API Key is correct. Error: {e}"


# -------- UI --------
st.title("Model 3: AI Resume Maker")
st.markdown("I saw your uploaded PDFs! Since programmatic conversion of static PDFs directly into editable Word docs breaks the layout severely, I have perfectly recreated 3 clean, minimalist, ATS-friendly templates from scratch directly in Python so you get a completely editable **Word (.docx)** and perfect **PDF**. Enter your details below to generate them!")

with st.sidebar:
    st.header("1. API Settings")
    api_key = st.text_input("Gemini API Key", type="password")
    
    st.header("3. Generate")
    template_choice = st.selectbox("Choose ATS Template", ["Classic ATS", "Modern Minimalist", "Executive"])
    st.info("Fill out your details on the right, get suggestions, and download!")

st.header("2. Your Information")

col1, col2 = st.columns(2)
with col1:
    name = st.text_input("Full Name", value="John Doe")
    email = st.text_input("Email Address", value="john@example.com")
with col2:
    phone = st.text_input("Phone Number", value="+1 234 567 8900")
    linkedin = st.text_input("LinkedIn Profile", value="linkedin.com/in/johndoe")

summary = st.text_area("Professional Summary", value="Dedicated professional with expertise in analyzing complex logic and delivering solutions...")
skills = st.text_area("Top Skills (comma separated)", value="Python, Data Analysis, Machine Learning, Leadership")
experience = st.text_area("Recent Experience (Use bullet points or dashes)", value="- Software Engineer at Tech Corp (2020-Present)\n  - Built APIs using Flask and FastAPI\n  - Improved system performance by 30%")
education = st.text_area("Education Details", value="- B.S. in Computer Science, State University\n- Graduated with honors")

data = {
    "name": name,
    "email": email,
    "phone": phone,
    "linkedin": linkedin,
    "summary": summary,
    "skills": skills,
    "experience": experience,
    "education": education
}

st.markdown("---")
st.subheader("Step 3: Review & Download")

col3, col4 = st.columns(2)

with col3:
    if st.button("Get AI Enhancements (Gemini)", use_container_width=True):
        if not api_key:
            st.error("Please insert your Gemini API key in the sidebar first!")
        else:
            with st.spinner("Gemini is analyzing your details..."):
                suggestions = get_resume_suggestions(api_key, data)
                st.success("AI Feedback Received:")
                st.write(suggestions)

with col4:
    # Generate files eagerly for download buttons
    docx_bytes = create_raw_docx(data, template_choice)
    try:
        pdf_bytes = create_raw_pdf(data, template_choice)
    except Exception as e:
        pdf_bytes = b""
        st.error(f"PDF Gen Error: {e}")
        
    st.download_button(
        label=f"📥 Download '{template_choice}' as Word (.docx)",
        data=docx_bytes,
        file_name="My_Resume.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
    st.download_button(
        label=f"📥 Download '{template_choice}' as PDF",
        data=pdf_bytes,
        file_name="My_Resume.pdf",
        mime="application/pdf",
        use_container_width=True
    )
